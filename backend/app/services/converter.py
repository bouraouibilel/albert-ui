import os
import re
import base64
import asyncio
import uuid
import pymupdf as fitz
import docx
import openpyxl
from bs4 import BeautifulSoup
import datetime
import time
from app.core.config import settings
from app.services.albert_client import albert_client

# Import résilient du logger
try:
    from app.core.logger import log_event
except ImportError:
    try:
        from core.logger import log_event
    except ImportError:
        def log_event(category: str, message: str, level: str = "INFO"):
            print(f"[{category}] {message}")

class DocumentConverter:
    """
    Service de conversion universelle de documents (PDF, DOCX, XLSX, TXT, HTML) vers Markdown (.md).
    Exclut les chaînes Base64 lourdes et génère des URLs d'images nettoyées et paramétrables pour le RAG.
    """

    @staticmethod
    def _sanitize_path_segment(name: str) -> str:
        """Nettoie une chaîne pour former un nom de dossier sûr pour le système de fichiers et les URLs."""
        clean = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', str(name or "")).strip('._-')
        return clean or "default"

    @staticmethod
    async def convert_to_markdown(file_path: str, filename: str, collection_name: str = "") -> dict:
        ext = os.path.splitext(filename)[1].lower()
        doc_prefix = str(uuid.uuid4())[:8]
        t0 = time.time()
        
        # Structure de répertoire pour les images : <collection_id_ou_nom>/<nom_document>
        col_folder = DocumentConverter._sanitize_path_segment(collection_name or "default")
        doc_base_name = os.path.splitext(filename)[0]
        doc_folder = DocumentConverter._sanitize_path_segment(doc_base_name)
        
        # Dossier physique : storage/images/<col_folder>/<doc_folder>/
        img_dir = os.path.join(settings.IMAGE_STORAGE_DIR, col_folder, doc_folder)
        os.makedirs(img_dir, exist_ok=True)
        
        # Préfixe d'URL : http://.../static/images/<col_folder>/<doc_folder>
        img_url_base = f"{settings.IMAGE_BASE_URL.rstrip('/')}/{col_folder}/{doc_folder}"
        
        log_event("CONVERTER", f"📄 Début de la conversion : '{filename}' (Format: {ext}, Collection: '{col_folder}', Prefix: {doc_prefix})")
        
        raw_text = ""
        tables_count = 0
        pages_count = 1

        if ext == ".pdf":
            raw_text, pages_count, tables_count = await DocumentConverter._convert_pdf(file_path, doc_prefix, img_dir, img_url_base)
        elif ext in [".docx", ".doc"]:
            raw_text, tables_count = await DocumentConverter._convert_docx(file_path, doc_prefix, img_dir, img_url_base)
        elif ext in [".xlsx", ".xls"]:
            raw_text, tables_count = DocumentConverter._convert_xlsx(file_path)
        elif ext in [".html", ".htm"]:
            raw_text = DocumentConverter._convert_html(file_path)
        elif ext in [".txt", ".md"]:
            log_event("CONVERTER", f"📖 Lecture directe du fichier texte '{filename}'...")
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
        else:
            log_event("CONVERTER", f"❌ Format non supporté: {ext}", level="ERROR")
            raise ValueError(f"Format non supporté: {ext}")

        cleaned_text = DocumentConverter._clean_markdown_text(raw_text)

        metadata_header = (
            "---\n"
            f"title: \"{filename}\"\n"
            f"converted_at: \"{datetime.datetime.now().isoformat()}\"\n"
            f"source_format: \"{ext}\"\n"
            f"pages: {pages_count}\n"
            f"tables_detected: {tables_count}\n"
            f"target_collection: \"{collection_name}\"\n"
            "---\n\n"
        )

        final_markdown = metadata_header + cleaned_text
        elapsed = round(time.time() - t0, 2)

        log_event("CONVERTER", f"✨ Conversion terminée pour '{filename}' en {elapsed}s | Pages: {pages_count}, Tableaux: {tables_count}, Caractères: {len(final_markdown)}")

        return {
            "markdown_content": final_markdown,
            "filename": filename,
            "pages_count": pages_count,
            "tables_count": tables_count,
            "char_count": len(final_markdown)
        }

    @staticmethod
    async def _convert_pdf(file_path: str, doc_prefix: str, img_dir: str, img_url_base: str) -> tuple[str, int, int]:
        """
        Conversion PDF avec sauvegarde physique des images dans storage/images/<col>/<doc>/ et URLs hiérarchisées.
        """
        doc = fitz.open(file_path)
        pages_count = len(doc)
        tables_count = 0
        md_content = []

        log_event("CONVERTER-PDF", f"🔍 Ouverture du PDF ({pages_count} pages)...")

        for page_num in range(pages_count):
            page = doc[page_num]
            md_content.append(f"\n\n## Page {page_num + 1}\n\n")
            
            tabs = page.find_tables()
            if tabs and tabs.tables:
                found_tab_len = len(tabs.tables)
                tables_count += found_tab_len
                log_event("CONVERTER-PDF", f"📊 Page {page_num + 1}: {found_tab_len} tableau(x) détecté(s)")

            page_text = page.get_text("text").strip()
            if page_text:
                md_content.append(f"{page_text}\n\n")

            image_list = page.get_images(full=True)
            diagram_count = 0
            for img_index, img_info in enumerate(image_list):
                if diagram_count >= 2:
                    break
                try:
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    img_ext = base_image["ext"]
                    
                    if len(image_bytes) < 15000:
                        continue

                    # Sauvegarde physique du fichier image sous storage/images/<col>/<doc>/
                    img_filename = f"img_p{page_num+1}_{img_index+1}.{img_ext}"
                    img_file_path = os.path.join(img_dir, img_filename)
                    with open(img_file_path, "wb") as f_img:
                        f_img.write(image_bytes)

                    # URL propre hiérarchisée
                    img_url = f"{img_url_base}/{img_filename}"
                    md_content.append(f"\n\n![Schéma Technique Page {page_num+1}]({img_url})\n\n")

                    # Analyse multimodale avec le modèle Vision 24B
                    img_b64 = base64.b64encode(image_bytes).decode("utf-8")
                    img_mime = f"image/{'jpeg' if img_ext.lower() in ['jpg', 'jpeg'] else img_ext.lower()}"
                    log_event("CONVERTER-PDF", f"🖼️ Image enregistrée : {img_filename} ({round(len(image_bytes)/1024, 1)} KB). URL: {img_url}")
                    
                    try:
                        vision_analysis = await asyncio.wait_for(
                            albert_client.describe_image(img_b64, mime_type=img_mime),
                            timeout=20.0
                        )
                        if vision_analysis:
                            log_event("CONVERTER-PDF", f"✅ Diagramme UML transcrit en Mermaid.js")
                            md_content.append(f"\n\n{vision_analysis}\n\n")
                    except asyncio.TimeoutError:
                        log_event("CONVERTER-PDF", f"⚠️ Timeout vision (20s) sur image {img_filename}", level="WARNING")
                    
                    diagram_count += 1
                except Exception as e:
                    log_event("CONVERTER-PDF", f"❌ Erreur image page {page_num + 1}: {e}", level="ERROR")

        return "".join(md_content), pages_count, tables_count

    @staticmethod
    async def _convert_docx(file_path: str, doc_prefix: str, img_dir: str, img_url_base: str) -> tuple[str, int]:
        """
        Conversion Word (.docx) avec sauvegarde des images dans storage/images/<col>/<doc>/ et URLs hiérarchisées.
        """
        doc = docx.Document(file_path)
        md_content = []
        tables_count = len(doc.tables)

        log_event("CONVERTER-WORD", f"🔍 Parcours du document Word ({len(doc.paragraphs)} paragraphes)...")
        img_counter = 0

        for element in doc.element.body:
            if element.tag.endswith('p'):
                p = docx.text.paragraph.Paragraph(element, doc)
                text = p.text.strip()
                if text:
                    style_name = p.style.name if p.style else ""
                    if style_name.startswith("Heading 1") or style_name.startswith("Titre 1"):
                        md_content.append(f"\n# {text}\n")
                    elif style_name.startswith("Heading 2") or style_name.startswith("Titre 2"):
                        md_content.append(f"\n## {text}\n")
                    elif style_name.startswith("Heading 3") or style_name.startswith("Titre 3"):
                        md_content.append(f"\n### {text}\n")
                    else:
                        md_content.append(f"\n{text}\n")

                for blip in element.xpath('.//a:blip'):
                    rId = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId and rId in doc.part.rels:
                        rel = doc.part.rels[rId]
                        if "image" in rel.target_ref:
                            try:
                                raw_ext = os.path.splitext(rel.target_ref)[1].lower().lstrip('.')
                                is_vector = raw_ext in ['emf', 'wmf', 'svg']
                                
                                image_blob = rel.target_part.blob
                                if len(image_blob) < 15000:
                                    continue
                                
                                img_ext = raw_ext if raw_ext in ['png', 'jpg', 'jpeg', 'webp', 'gif', 'bmp'] else 'png'
                                img_mime = f"image/{'jpeg' if img_ext in ['jpg', 'jpeg'] else img_ext}"
                                
                                img_counter += 1
                                img_filename = f"img_{img_counter}.{img_ext}"
                                img_file_path = os.path.join(img_dir, img_filename)
                                with open(img_file_path, "wb") as f_img:
                                    f_img.write(image_blob)

                                img_url = f"{img_url_base}/{img_filename}"
                                md_content.append(f"\n\n![Schéma Technique Word]({img_url})\n\n")

                                if not is_vector:
                                    img_b64 = base64.b64encode(image_blob).decode("utf-8")
                                    log_event("CONVERTER-WORD", f"🖼️ Image Word enregistrée : {img_filename} ({round(len(image_blob)/1024, 1)} KB). URL: {img_url}")
                                    
                                    try:
                                        vision_analysis = await asyncio.wait_for(
                                            albert_client.describe_image(img_b64, mime_type=img_mime),
                                            timeout=20.0
                                        )
                                        if vision_analysis:
                                            log_event("CONVERTER-WORD", f"✅ Diagramme UML Word transcrit en Mermaid.js")
                                            md_content.append(f"\n\n{vision_analysis}\n\n")
                                    except asyncio.TimeoutError:
                                        log_event("CONVERTER-WORD", f"⚠️ Timeout vision (20s) sur image Word {img_filename}", level="WARNING")
                                else:
                                    log_event("CONVERTER-WORD", f"🖼️ Image vectorielle ({raw_ext.upper()}) enregistrée sans appel LLM Vision : {img_filename}")
                            except Exception as e:
                                log_event("CONVERTER-WORD", f"❌ Erreur image Word: {e}", level="ERROR")

            elif element.tag.endswith('tbl'):
                table = docx.table.Table(element, doc)
                md_content.append("\n\n")
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    rows.append(cells)
                if rows and len(rows) > 0:
                    header = rows[0]
                    md_content.append("| " + " | ".join(header) + " |\n")
                    md_content.append("| " + " | ".join(["---"] * len(header)) + " |\n")
                    for r in rows[1:]:
                        md_content.append("| " + " | ".join(r) + " |\n")
                md_content.append("\n\n")

        return "".join(md_content), tables_count

    @staticmethod
    def _convert_xlsx(file_path: str) -> tuple[str, int]:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        md_content = []
        tables_count = len(wb.sheetnames)

        log_event("CONVERTER-EXCEL", f"🔍 Traitement des {tables_count} feuille(s) Excel...")

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            md_content.append(f"\n# Feuille: {sheet_name}\n\n")
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            
            clean_rows = []
            for row in rows:
                if any(cell is not None for cell in row):
                    clean_rows.append([str(c or "").strip().replace("\n", " ") for c in row])

            if clean_rows:
                header = clean_rows[0]
                md_content.append("| " + " | ".join(header) + " |\n")
                md_content.append("| " + " | ".join(["---"] * len(header)) + " |\n")
                for r in clean_rows[1:]:
                    md_content.append("| " + " | ".join(r) + " |\n")
            md_content.append("\n\n")

        return "".join(md_content), tables_count

    @staticmethod
    def _convert_html(file_path: str) -> str:
        log_event("CONVERTER-HTML", f"🔍 Extraction du contenu HTML...")
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
        
        for element in soup(["script", "style", "nav", "footer"]):
            element.decompose()

        md_content = []
        for elem in soup.find_all(["h1", "h2", "h3", "h4", "p", "table"]):
            if elem.name in ["h1", "h2", "h3", "h4"]:
                level = int(elem.name[1])
                md_content.append(f"\n{'#' * level} {elem.get_text().strip()}\n")
            elif elem.name == "p":
                text = elem.get_text().strip()
                if text:
                    md_content.append(f"\n{text}\n")
            elif elem.name == "table":
                rows = []
                for tr in elem.find_all("tr"):
                    cells = [td.get_text().strip().replace("\n", " ") for td in tr.find_all(["td", "th"])]
                    if cells:
                        rows.append(cells)
                if rows:
                    md_content.append("\n\n| " + " | ".join(rows[0]) + " |\n")
                    md_content.append("| " + " | ".join(["---"] * len(rows[0])) + " |\n")
                    for r in rows[1:]:
                        md_content.append("| " + " | ".join(r) + " |\n")
                    md_content.append("\n\n")

        return "".join(md_content)

    @staticmethod
    def _clean_markdown_text(text: str) -> str:
        lines = text.splitlines()
        cleaned_lines = []
        empty_count = 0
        for line in lines:
            if not line.strip():
                empty_count += 1
                if empty_count <= 2:
                    cleaned_lines.append("")
            else:
                empty_count = 0
                cleaned_lines.append(line)
        return "\n".join(cleaned_lines)
